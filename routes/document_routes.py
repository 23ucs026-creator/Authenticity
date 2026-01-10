import os
from flask import Blueprint, request, jsonify, current_app
from flask_jwt_extended import jwt_required, get_jwt_identity
from werkzeug.utils import secure_filename
from extensions import db
from models.document import Document

# Text extraction functions
def extract_pdf_text(file_path):
    from pdfminer.high_level import extract_text
    return extract_text(file_path)  # Simple extraction [web:31]

def extract_docx_text(file_path):
    from docx import Document
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_txt_text(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

document_bp = Blueprint("documents", __name__)
ALLOWED_EXT = {"pdf", "docx", "txt"}


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXT


@document_bp.route("/documents/upload", methods=["POST"])
@jwt_required()
def upload_document():
    # 1. Check file part
    if "file" not in request.files:
        return jsonify({"msg": "No file part"}), 400

    file = request.files["file"]

    if file.filename == "":
        return jsonify({"msg": "No selected file"}), 400

    if not allowed_file(file.filename):
        return jsonify({"msg": "Unsupported file type"}), 400

    # 2. Save file to uploads folder
    filename = secure_filename(file.filename)
    upload_folder = os.path.join(current_app.root_path, "uploads")
    os.makedirs(upload_folder, exist_ok=True)
    path = os.path.join(upload_folder, filename)
    file.save(path)

<<<<<<< HEAD
    # Extract text based on file type
    ext = filename.rsplit(".", 1)[1].lower()
    try:
        if ext == 'pdf':
            original_text = extract_pdf_text(path)
        elif ext == 'docx':
            original_text = extract_docx_text(path)
        elif ext == 'txt':
            original_text = extract_txt_text(path)
        else:
            original_text = None
    except Exception as e:
        original_text = f"Extraction failed: {str(e)}"

    user_id = get_jwt_identity()
=======
    # 3. Get user id from JWT (stored as string, convert to int)
    user_id = int(get_jwt_identity())

    ext = filename.rsplit(".", 1)[1].lower()

    # 4. Create Document record (text extraction can be added later)
>>>>>>> 0c8b76ff396c77d759de0186bd270e7333d31c59
    doc = Document(
        user_id=user_id,
        filename=filename,
        file_type=ext,
        original_text=original_text,
    )
    db.session.add(doc)
    db.session.commit()

    return jsonify({"msg": "File uploaded and text extracted", "document_id": doc.id, "text_length": len(original_text or "")}), 201
